from pathlib import Path
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement

def insert_break_before(paragraph):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type','page')
    r.append(br); p.append(r)
    paragraph._p.addprevious(p)

root=Path('/mnt/data/M2-computer-vision-technologies/kim')
for path in root.glob('*/student_guide.docx'):
    doc=Document(path)
    for p in doc.paragraphs:
        if 'Рекомендуемая структура README' in p.text:
            insert_break_before(p); break
    doc.save(path)
for path in root.glob('*/teacher_rubric.docx'):
    doc=Document(path)
    for p in doc.paragraphs:
        if p.text.strip().startswith('4. Градации качества'):
            insert_break_before(p); break
    doc.save(path)
print('patched')
